#coding=utf-8

'''模块说明：
本模块实现的功能，主要是将分析结果的文本文件，按你的自定义标签，塞到你自定义的报告模板对应标签位置当中。

经典代码如下：
import docx # 如果提示没有这个模块 pip install python-docx
resultFilePath = 'result.txt' #可逐行读取的分析结果文件，不限于txt
templateFilePath = '模版.docx' #自定义模板文件
saveFilePath = '2.docx' #输出路径，只写文件名就是程序当前目录下
report = docx.Document(unicode(templateFilePath, 'utf-8')) #依据模板生成一份报告
report = fillAnalyseResultFile(resultFilePath,report) #填充一个结果文件
report = fillAnalyseResultFile(resultFilePath,report) #可连续填充多个结果文件
resultMap = {'#[NAME]#':'徐超','#[SEX]#':'男'}  #可不基于文件，而是直接将内存中的分析结果，集中到一个字典中，再将字典写入模板文件
report = fillAnalyseResultMap(resultMap,report)
report.save(saveFilePath) #最后记得保存，不然一场空哦

主要的目的，是为了实现分析过程与模板样式之间解开耦合，所有关于段落、字体、风格的调整，都不需要修改任何代码，
尽量让项目负责人自己根据客户要求调整模板，不需要专业编程人员协助。
可参见result.txt中的示例，标签可以自定义，但需要遵守的约定是#[XXXXX]#=内容，XXXXX可替换成任何内容
等号后紧接内容，如果有回车空格，会认为是正常内容，所以如果内容是一个表格，等号紧接第一行开始的字符
内容可以是多行的，不过多行的内容写进去会保持原样，也是多行的

内容是表格时的特殊约定：
1、标签统一使用“#[FILLTABLE-1-4(0,1|7)]#=”这种形式，其中不可修改的是“#[FILLTABLE]#=”这些字符，“-1-4”这几个字符可以任意替换增删，
以区别别的标签，“(0,1|7)”这几个字符是表示需要合并相同值的列，|左边是以整行完整考虑合并的列，例如第三列相同值但第二列不同，
则第三列不会合并，按实际情况看应该从0开始，多列以半角逗号隔开，|右边则是单列考虑合并的列，按实际考虑一般都是最后一列
表格的填充，是从标签所在行和单元格往下填充，右侧超过内容将丢失，下方超过表格将自动补行，所以应该让模板保留一行，标签替换保留行第一个单元格
并且为了调整和固定样式，每个单元格都应该有旧的内容，手动调整好格式后，填充时内容将被替换，格式将被保留，后面新生成的行的格式将与保留的行一致
2、表格以回车识别行，所以单元格内容是多行的时候，需要用三个型号代表换行，如果使用了换行则认为出现一个不规范的列，生成过程可能会报错；
三个星号换行的规则，在其他地方也适用。
3、表格不要采用续表的形式拆开，而是只留一个表头和一个内容行，内容行第一格填上表格的标签，后面复制一些旧的内容进去固定格式，然后在右键的表格属性中，
勾选行属性页中的“在各页顶端以标题行形式重复出现”的复选框，因为表格的行数是不定的，就算你安排好表格位置，插入内容后也会挤到后面去。

word-writer升级说明：20170628
本次增加插入图片功能
标签举例：#[IMAGE2(33,11)]#
#[IMAGE是插入图片特有的，小括号里面是宽和高，单位是“十万dpi”（比word边上的标尺小一点，宽33高11对应的标尺大概是22和8），整个小括号部分可省略，图片将按原大小插入
值举例：a.jpg
就是图片的路径，可以是相对路径或绝对路径


注意的问题：
1、想要对分析逻辑、输出结果的修改，仍然需要修改的代码或数据库，这只是提供一个word版本报告的输出。
2、由于docx包固有的一些bug，单元格的字体无法获取，如果更换模板，注意在代码中人工指定单元格中文字的字体，此处表格统一使用的字体为'Calibri'，
与当前和谐模板基本一致20170525。
3、标签的使用应该尽量统一，否则某些格式设置可能考虑不到；
4、与现在的分析结果需要做一个对应过程才能使用，模板和对应过程需要每个项目特殊定制，不过这个过程是一次制作重复使用，开始可以做一做；
后期少量修改标签对应方案总比修改html样式代码轻松。
5、需要适当的多线程优化，提高替换效率；但是经过测试多线程后比单线程反而慢很多，原因不明，所以目前仍然采用单线程写入20170526；
6、表格中，风险高风险低这种局部样式的修改这个插件也不能很好地支持，目前只能整段的修改，也就是说一个单元格“风险较高”可以整体变成橙色，
但是把“较高”两个字单独变成橙色暂时比较困难;
7、由于word文档自身深层次结构的原因，建议将标签写好之后整体粘贴到模板中，然后整体调整格式，
否则有可能导致标签不识别或者替换后仍然残留部分标签的奇怪情况发生；如果发现奇怪情况，将标签原地剪切粘贴，选择仅保留文本，然后整体调整格式。
还有，编辑模板请关闭word的语法和拼写检查
'''

#创建 Document 对象，相当于在内存中拷贝模板新建一个 word 文档

import docx
# import threading
# import datetime
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

specialRowMerge = ['药物超敏反应预测','药物敏感性预测','药物体内代谢情况','药物毒副反应预测']
#colorRules = {'风险高':RGBColor(0xFF, 0x00, 0x00),'风险较高':RGBColor(0xFF, 0xA5, 0x00),'风险低':RGBColor(0x00, 0xFF, 0x00)}
colorRules = {}


def __color__(paragraph, colorRules):
    # paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    for key in colorRules:
        alltext = paragraph.text
        if alltext.__contains__(unicode(key,'utf-8')):
            paragraph.runs[0].font.color.rgb = colorRules[key]
            # b = paragraph.runs[0].bold
            # n = paragraph.runs[0].font.name
            # s = paragraph.runs[0].font.size
            # runList = paragraph.runs
            # for run in runList:
            #     if str(paragraph.text).__contains__(key):
            #         splits = str(run.text).split(key)
            #         run.text = ''
            #         lastColorRun = None
            #         for pf in splits:
            #             normalRun = paragraph.add_run(pf,run.style)
            #             colorRun = paragraph.add_run(key, run.style)
            #             colorRun.font.color.rgb = colorRules[key]
            #             lastColorRun = colorRun
            #         #run.font.color.rgb = colorRules[key]
            #         lastColorRun.text = ''

def __addPicture__(document, searchString, picturePath):
    for paragraph in document.paragraphs:
        textAll = paragraph.text
        if textAll.__contains__(searchString):
            __replaceParagraph__(paragraph,searchString," ")
            firstRun = paragraph.runs[0]

            if str(searchString).__contains__('(') and str(searchString).__contains__(')'):
                L = searchString.find('(')+1
                R = searchString.find(')')
                Str = searchString[L:R]
                wh = str(Str).split(',')
                width = int(wh[0].strip())
                height = int(wh[1].strip())
                firstRun.add_picture(picturePath,width*100000,height*100000)
            else:
                firstRun.add_picture(picturePath)


def __replaceParagraph__(paragraph,searchString,replaceString):
    if searchString=='#[REFERENCE]#':#参考文献的格式需要分段后挨个设置
        pList =replaceString.split('\n')
        replaceString = ' '
        style = paragraph.style
        alignment = paragraph.alignment
        paragraph_format = paragraph.paragraph_format
        b = paragraph.runs[0].bold
        n = paragraph.runs[0].font.name
        s = paragraph.runs[0].font.size
        while(len(pList)>0):
            content = pList.pop(0)
            newParagraph = paragraph.insert_paragraph_before(unicode(content, 'utf-8'),style)
            newParagraph.alignment =alignment
            #newParagraph.paragraph_format = paragraph_format
            __setParagraphformat__(newParagraph,paragraph_format)
            newParagraph.runs[0].bold =b
            newParagraph.runs[0].font.name=n
            newParagraph.runs[0].font.size=s

    for run in paragraph.runs:
        text = run.text
        if text.__contains__(searchString):
            text = text.replace(searchString, replaceString)
            run.text = unicode(text.replace('***','\n'), 'utf-8')

def __setParagraphformat__(newParagraph, paragraph_format):
    newParagraph.paragraph_format.alignment = paragraph_format.alignment
    newParagraph.paragraph_format.first_line_indent=paragraph_format.first_line_indent
    newParagraph.paragraph_format.keep_together=paragraph_format.keep_together
    newParagraph.paragraph_format.keep_with_next=paragraph_format.keep_with_next
    newParagraph.paragraph_format.keep_with_next=paragraph_format.keep_with_next
    newParagraph.paragraph_format.line_spacing=paragraph_format.line_spacing
    newParagraph.paragraph_format.line_spacing_rule=paragraph_format.line_spacing_rule
    newParagraph.paragraph_format.page_break_before=paragraph_format.page_break_before
    newParagraph.paragraph_format.right_indent=paragraph_format.right_indent
    newParagraph.paragraph_format.space_after=paragraph_format.space_after
    newParagraph.paragraph_format.space_before=paragraph_format.space_before
    #newParagraph.paragraph_format.tab_stops=paragraph_format.tab_stops
    newParagraph.paragraph_format.widow_control=paragraph_format.widow_control


def __setCellCenter__(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), "center")
    tcPr.append(tcVAlign)

def deleteEmptyTable(document,tagString):
    tableList = []
    for table in document.tables:
        skip = False
        checkrowMax = 3
        rowidx = 1
        for row in table.rows:
            if rowidx>checkrowMax:
                break
            else:
                rowidx+=1
            if skip:break

            checkcolMax = 3
            colidx = 1
            for cell in row.cells:
                if colidx > checkcolMax:
                    break
                else:
                    colidx += 1
                if skip: break
                textAll = cell.text
                if textAll.__contains__(tagString):
                    if table not in tableList:
                        tableList.append(table)
                        skip = True
    for table in tableList:
        for row in table.rows:
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)
    return document

def __fillTable__(document,searchString,replaceString,mergeColumnsWithSingleForce,specialRowMerge):
    startFilling =False
    fillTableIdx = 0
    fillRowIdx =0
    fillCellIdx = 0
    rowsContents = str(replaceString).split('\n')#如果单元格中需要换行，请用三个星号代替：***
    if searchString == '#[FILLTABLE-Medolanzapine]#':
        pass
    rowContentIdx = 0

    # 先定位开始填充的位置，标记在哪个表格，行，单元格
    # 在有标记的单元格开始填充，一般标记应该放在表格第一行的第一个单元格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                textAll = cell.text
                if textAll.__contains__(searchString):
                    startFilling = True
                if (startFilling):
                    break
                fillCellIdx+=1
            if (startFilling):
                break
            fillRowIdx+=1
            fillCellIdx=0
        if (startFilling):
            break
        fillTableIdx+=1
        fillRowIdx=0

    #开始填充 开始单元格的坐标不变 右边超出的不填入 行不够的生成新的行
    #后期需要合并相同值的单元格
    if startFilling:
        table = document.tables[fillTableIdx]
        rowSkip = fillRowIdx
        rowForStyle = None;
        for row in table.rows:
            if rowSkip>0:
                rowSkip -=1
                continue
            rowForStyle = table.rows[1]#格式参考第一行，需要保证模板每个表格至少有一个内容行，可以有内容并设定格式，内容将覆盖格式将保留

            #表格预留的空行比数据还多
            if rowContentIdx+1 > len(rowsContents):
                break

            cellsContent = rowsContents[rowContentIdx].split('\t')

            if len(cellsContent)<2:
                continue

            rowContentIdx+=1
            cellContentIdx = 0
            cellSkip = fillCellIdx
            for cell in row.cells:
                if cellSkip > 0:
                    cellSkip -= 1
                    continue
                if len(cell.paragraphs[0].runs)>0:
                    cell.paragraphs[0].runs[0].text = unicode(cellsContent[cellContentIdx].replace('***','\n'), 'utf-8')

                elif len(cell.paragraphs)>0:
                    cell.paragraphs[0].text = unicode(cellsContent[cellContentIdx].replace('***','\n'), 'utf-8')
                else:
                    cell.text = unicode(cellsContent[cellContentIdx].replace('***','\n'), 'utf-8')
                __setCellCenter__(cell)
                __color__(cell.paragraphs[0], colorRules)
                cellContentIdx+=1

        #这里都是补的单元格
        while rowContentIdx<len(rowsContents):
            cellsContent = rowsContents[rowContentIdx].split('\t')
            if len(cellsContent)<2:
                continue
            rowContentIdx += 1
            cellContentIdx = 0


            #格式 参照表格第一行设置
            newRow = table.add_row()
            cellStyleNumber=0

            cellSkip = fillCellIdx
            for cell in newRow.cells:
                if cellSkip > 0:
                    cellSkip -= 1
                    cellStyleNumber+=1
                    continue
                referCell = rowForStyle.cells[cellStyleNumber]
                cellStyleNumber += 1

                # try:
                cell.text = unicode(cellsContent[cellContentIdx].replace('***','\n'), 'utf-8')
                # except:
                #     print searchString
                #     print cellContentIdx
                #     print table.rows[0].cells[0].text
                #     return
                cell.paragraphs[0].style = referCell.paragraphs[0].style
                cell.paragraphs[0].alignment = referCell.paragraphs[0].alignment
                cell.paragraphs[0].runs[0].bold = referCell.paragraphs[0].runs[0].bold
                cell.paragraphs[0].runs[0].font.name = 'Calibri'#referCell.paragraphs[0].runs[0].font.name #可能是包的bug（中文的字体名），字体名称无法正确获得
                cell.paragraphs[0].runs[0].font.size = referCell.paragraphs[0].runs[0].font.size

                __color__(cell.paragraphs[0], colorRules)
                __setCellCenter__(cell)

                cellContentIdx += 1

        #整个表格的内容都已经填充完毕
        #在列方向对相同值的单元格进行合并，而且只有前一列合并过的组，后面才有合并可能，并不是一旦相同两个相邻单元格就合并
        if len(mergeColumnsWithSingleForce)>0:
            #整个表格逐列递归，找出所有应该合并的相邻单元格
            shouleMergesMap = {}#整个表中，需要所有列，所有组，所有应该合并的表格
            columnIdx = 0
            for column in table.columns:
                columnToCheck = column
                formerColumnGroup = []
                if columnIdx >0:
                    formerColumnGroup = shouleMergesMap[columnIdx-1]
                shouleMergesMap[columnIdx] = __checkMerge__(columnIdx,columnToCheck,formerColumnGroup)
                columnIdx+=1

            #合并指定列的，上面找出应该合并的单元格
            mergeColumns = mergeColumnsWithSingleForce.pop(0)
            for force in mergeColumnsWithSingleForce:
                shouleMergesMap[force] = __forceMergeSingleColumn__(table.columns[force])
                mergeColumns.append(force)

            for colIdx in mergeColumns:
                groups = shouleMergesMap[colIdx]
                for group in groups:
                    firstCell = None
                    if len(group)>1:#两个及以上应该合并 则合并
                        cells = table.columns[colIdx].cells
                        cellList = []
                        for cellIdx in group:
                            cellList.append(cells[cellIdx])
                        firstCell = cellList.pop(0)
                        for xxx in cellList:
                            __mergeCell__(firstCell, xxx)

            for row in table.rows:
                cells = row.cells
                firstCell = cells[0]
                cellList = []
                stop = False
                for xxx in cells:
                    fStr = firstCell.text.encode("utf-8").strip()
                    xStr = xxx.text.encode("utf-8").strip()
                    if fStr == xStr and xStr in specialRowMerge:
                        cellList.append(xxx)
                        stop = True #加入过单元格的 参照单元格不再移动 所以每一行只允许一次行方向连续区域的特别合并
                    else:
                        if stop:
                            break
                        firstCell = xxx
                if len(cellList)>1:
                    cellList.pop(0)
                    for xxx in cellList:
                        __mergeCell__(firstCell,xxx)

def __mergeCell__(firstCell,secondCell):
    st = firstCell.paragraphs[0].style
    a = firstCell.paragraphs[0].alignment
    b = firstCell.paragraphs[0].runs[0].bold
    n = firstCell.paragraphs[0].runs[0].font.name
    s = firstCell.paragraphs[0].runs[0].font.size
    firstCell.text = ''
    firstCell.merge(secondCell)
    firstCell.text = firstCell.text.strip().replace('***', '\n')
    firstCell.paragraphs[0].style = st
    firstCell.paragraphs[0].alignment = a
    firstCell.paragraphs[0].runs[0].bold = b
    firstCell.paragraphs[0].runs[0].font.name = n
    firstCell.paragraphs[0].runs[0].font.size = s
    __setCellCenter__(firstCell)

def __check__(columnToCheck,start,end):
    column = columnToCheck
    groupsListTemp = []#列表的列表

    thisCellIdx = start
    nextcellIdx = start+1
    breakFlag = False
    while thisCellIdx < end:# 夹子 跳跃 检查方法
        if breakFlag:
            break
        mergeCellsList = []
        mergeCellsList.append(thisCellIdx)
        while nextcellIdx <= end:
            thisCellText = column.cells[thisCellIdx].text.strip()
            nextCellText = column.cells[nextcellIdx].text.strip()
            if (thisCellText == nextCellText):
                mergeCellsList.append(nextcellIdx)
                nextcellIdx+=1
            else:
                groupsListTemp.append(mergeCellsList)
                thisCellIdx = nextcellIdx
                nextcellIdx = nextcellIdx+1
                mergeCellsList = []
                mergeCellsList.append(thisCellIdx)
            if nextcellIdx > end:#夹子后腿移出指定区域
                breakFlag = True
                groupsListTemp.append(mergeCellsList)

    return groupsListTemp

def __forceMergeSingleColumn__(columnToCheck):
    groups = []
    # 唯一一列，从头到尾扫描
    start = 0
    end = len(columnToCheck.cells) - 1
    groupsListTemp = __check__(columnToCheck, start, end)
    if len(groupsListTemp) > 0:
        for singleGroup in groupsListTemp:
            groups.append(singleGroup)

    return groups

def __checkMerge__(columnIdx,columnToCheck,formerColumnGroup):
    groups = []

    if columnIdx==0:#第一列，从头到尾扫描
        start = 0
        end = len(columnToCheck.cells)-1
        groupsListTemp = __check__(columnToCheck,start,end)
        if len(groupsListTemp) > 0:
            for singleGroup in groupsListTemp:
                groups.append(singleGroup)
    else:#其他列，按照分组扫描，但是分组内只有一个元素的，跳过
        for singleFormerGroup in formerColumnGroup:
            if len(singleFormerGroup) <2:
                continue
            else:
                start = singleFormerGroup[0]
                end = singleFormerGroup[len(singleFormerGroup)-1]
                groupsListTemp = __check__(columnToCheck, start, end)
            if len(groupsListTemp)>0:
                for singleGroup in groupsListTemp:
                    groups.append(singleGroup)

    return groups

#需要改成多线程提高速度
def __replaceAll__(document,searchString,replaceString):
    for paragraph in document.paragraphs:
        textAll = paragraph.text
        if textAll.__contains__(searchString):
            __replaceParagraph__(paragraph,searchString,replaceString)


    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                textAll = cell.text
                if textAll.__contains__(searchString):
                    for paragraph in cell.paragraphs:
                        __replaceParagraph__(paragraph, searchString, replaceString)

def fillAnalyseResultFile(resultFilePath,report):
    print 'filling analyse result of ', resultFilePath
    # st = datetime.datetime.now()
    resultFile = open(resultFilePath)
    resultMap = {}#为了支持多行，需要先全部读入
    lastSearchString = ''
    for line in resultFile:
        if line.startswith('#[') and line.__contains__(']#='):#如果这一行不以标签开头，则归入上一行的结果
            ind = line.find('=')
            searchString = line[0:ind]#标签
            lastSearchString=searchString
            replaceString = line[ind + 1:]#替换的内容，只是当前行，可能后面还有的
            resultMap[searchString] = replaceString
        else:
            resultMap[lastSearchString] += line

    report = fillAnalyseResultMap(resultMap,report)

    print 'finish filling ',resultFilePath
    resultFile.close()
    # se = datetime.datetime.now()
    # print 'elapse seconds, ',(se-st).seconds
    return report

def fillAnalyseResultMap(resultMap,report):
    print 'filling analyse result of resultMap, about 1 minute'
    # threadList = [] #改成多线程，同时写入多个结果
    # 经过测试，同一个填充任务，多线程耗时19秒，单线程耗时10秒，原因不明，放弃使用多线程

    for key in resultMap:
        searchString=key
        print 'filling content of tag: ',searchString
        replaceString=str(resultMap[key]).strip()
        if str(searchString).__contains__('#[FILLTABLE'):#如果要指定合并相同值的列，用半角括号半角逗号分隔如(0,7)
            mergeColumns = []# #[FILLTABLE-1-4(0,1|7)]# 表示0,1列正常排序（必须从0开始），|后面的列单独排序合并
            mergeColumnsList = []
            if str(searchString).__contains__('(') and str(searchString).__contains__(')'):
                mergeL = searchString.find('(')+1
                mergeR = searchString.find(')')
                mergeStr = searchString[mergeL:mergeR]
                mergesP = str(mergeStr).split('|')
                merges = mergesP.pop(0).split(',')
                for merg in merges:
                    mergeColumns.append(int(merg))
                mergeColumnsList.append(mergeColumns)
                for singleM in mergesP:
                    mergeColumnsList.append(int(singleM))
            __fillTable__(report, searchString, replaceString,mergeColumnsList,specialRowMerge)
            # threadList.append(threading.Thread(target=__fillTable__, args=(report, searchString, replaceString,mergeColumnsList)))
        elif str(searchString).__contains__('#[IMAGE'):
            __addPicture__(report, searchString, replaceString)
        else:
            __replaceAll__(report, searchString, replaceString)
            # threadList.append(threading.Thread(target=__replaceAll__, args=(report, searchString, replaceString)))

    # for t in threadList:
    #     t.setDaemon(True)
    #     t.start()
    #
    # for t in threadList:
    #     t.join()

    return report

#########################################  测试代码  ####################################
if __name__ == "__main__": 
	resultFilePath = 'result.txt'
	templateFilePath = '模版.docx'
	saveFilePath = '2.docx'
	report = docx.Document(unicode(templateFilePath, 'utf-8'))
	report = fillAnalyseResultFile(resultFilePath,report)
	report.save(saveFilePath)



